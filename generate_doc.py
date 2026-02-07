"""
Generate Word Document from Project Documentation
Converts PROJECT_DOCUMENTATION.md to a properly formatted Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path


def parse_markdown(md_content):
    """Parse markdown content into sections."""
    lines = md_content.split('\n')
    sections = []
    current_section = None
    current_content = []
    in_code_block = False
    code_block_content = []
    
    for line in lines:
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                if current_section:
                    current_content.append(('code', '\n'.join(code_block_content)))
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # Handle headers
        if line.startswith('# '):
            if current_section:
                sections.append((current_section, current_content))
            current_section = ('h1', line[2:].strip())
            current_content = []
        elif line.startswith('## '):
            if current_section:
                sections.append((current_section, current_content))
            current_section = ('h2', line[3:].strip())
            current_content = []
        elif line.startswith('### '):
            current_content.append(('h3', line[4:].strip()))
        elif line.startswith('#### '):
            current_content.append(('h4', line[5:].strip()))
        elif line.startswith('- ') or line.startswith('* '):
            current_content.append(('bullet', line[2:].strip()))
        elif line.startswith('|'):
            current_content.append(('table_row', line))
        elif line.strip():
            current_content.append(('para', line.strip()))
    
    if current_section:
        sections.append((current_section, current_content))
    
    return sections


def create_word_document(md_path, output_path):
    """Create Word document from markdown file."""
    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 76, 153)
    
    # Parse markdown
    sections = parse_markdown(md_content)
    
    # Add content
    for section_header, content in sections:
        header_type, header_text = section_header
        
        # Add section header
        if header_type == 'h1':
            # Check if it's the main title
            if 'Trustworthy RAG Agent' in header_text:
                p = doc.add_paragraph(header_text, style='Title')
            else:
                doc.add_heading(header_text, level=1)
        elif header_type == 'h2':
            doc.add_heading(header_text, level=2)
        
        # Add section content
        for item_type, item_content in content:
            if item_type == 'h3':
                doc.add_heading(item_content, level=3)
            elif item_type == 'h4':
                p = doc.add_paragraph()
                run = p.add_run(item_content)
                run.bold = True
                run.font.size = Pt(11)
            elif item_type == 'bullet':
                # Clean markdown formatting
                clean_text = item_content.replace('**', '').replace('`', '')
                doc.add_paragraph(clean_text, style='List Bullet')
            elif item_type == 'code':
                # Add code block with formatting
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(item_content)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            elif item_type == 'table_row':
                # Skip table formatting for now (complex)
                pass
            elif item_type == 'para':
                # Clean markdown formatting
                clean_text = item_content
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)  # Bold
                clean_text = re.sub(r'`(.*?)`', r'\1', clean_text)  # Code
                clean_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_text)  # Links
                
                if clean_text.strip() and clean_text.strip() != '---':
                    p = doc.add_paragraph(clean_text)
    
    # Save document
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    project_dir = Path(__file__).parent
    md_path = project_dir / "PROJECT_DOCUMENTATION.md"
    output_path = project_dir / "Project_Documentation.docx"
    
    create_word_document(md_path, output_path)
